from django.shortcuts import render, redirect
from django.http import HttpResponse
from .mongodb import save_to_mongodb
from docx import Document
import re
from django.contrib.auth.decorators import login_required
from .firebase_utils import AssignmentManager
from django.http import JsonResponse

def submissions(request):
    print("submissions")
    if request.method == 'POST' and request.FILES['file']:
        uploaded_file = request.FILES['file']

        assignment_id = request.session.get('current_assignment_id') or request.POST.get('assignment_id')
        if not assignment_id:
            return HttpResponse('No assignment selected', status=400)

        # Check if the uploaded file is a .docx file
        if not uploaded_file.name.endswith('.docx'):
            return HttpResponse('Invalid file type. Please upload a .docx file.', status=400)

        # Read the .docx file
        try:
            document = Document(uploaded_file)
            file_content = "\n".join([para.text for para in document.paragraphs])
        except Exception as e:
            print(f"Error reading .docx file: {e}")
            return HttpResponse('Failed to read .docx file', status=400)

        print(f"File content: {file_content[:100]}...")  # Log the first 100 characters of the file content

        # Parse the content

        try:
            name_match = re.search(r'Name:\s*(.*)', file_content)
            class_match = re.search(r'Class:\s*(.*)', file_content)
            sub_match = re.search(r'Subject:\s*(.*)', file_content)
            qa_matches = re.findall(r'Question:\s*(.*?)\s*Answer:\s*(.*?)(?=\s*Question:|\s*$)', file_content, re.DOTALL)

            if not name_match or not class_match or not qa_matches:
                return HttpResponse('Invalid content format in .docx file', status=400)

            name = name_match.group(1).strip()
            class_name = class_match.group(1).strip()
            sub = sub_match.group(1).strip()
            questions_answers = []
            for question, answer in qa_matches:
                questions_answers.append({
                    'question': question.strip(),
                    'answer': answer.strip()
                })

            print("data extracted")
        except Exception as e:
            print(f"Error parsing content: {e}")
            return HttpResponse('Failed to parse content in .docx file', status=400)

        # Save to Firebase
        assignment_manager = AssignmentManager()
        if assignment_manager.save_submission(name, class_name, questions_answers, sub, assignment_id, username=request.user.username):
            return HttpResponse('File uploaded successfully')
        else:
            return HttpResponse('Failed to save submission', status=500)

    return render(request, '../static/templates/submissions/submission_portal.html')

def submissions_backup(request):
    print("submissions")
    if request.method == 'POST' and request.FILES['file']:
        # Check if user is authenticated
        if not request.user.is_authenticated:
            return HttpResponse('User not authenticated', status=401)

        # Check if assignment_id is in the session or POST data
        assignment_id = request.session.get('current_assignment_id') or request.POST.get('assignment_id')
        if not assignment_id:
            return HttpResponse('No assignment selected', status=400)

        uploaded_file = request.FILES['file']

        # Check if the uploaded file is a .docx file
        if not uploaded_file.name.endswith('.docx'):
            return HttpResponse('Invalid file type. Please upload a .docx file.', status=400)

        # Read the .docx file
        try:
            document = Document(uploaded_file)
            file_content = "\n".join([para.text for para in document.paragraphs])
        except Exception as e:
            print(f"Error reading .docx file: {e}")
            return HttpResponse('Failed to read .docx file', status=400)

        # Parse the answers from the document
        try:
            # Find all Answer sections
            qa_matches = re.findall(r'Answer:\s*(.*?)(?=\s*Question:|\s*$)', file_content, re.DOTALL)

            if not qa_matches:
                return HttpResponse('No answers found in the document', status=400)

            # Prepare questions_answers list with just answers
            questions_answers = [
                {
                    'answer': answer.strip()
                } for answer in qa_matches
            ]

            print("data extracted")
        except Exception as e:
            print(f"Error parsing content: {e}")
            return HttpResponse('Failed to parse content in .docx file', status=400)

        # Save to Firebase
        assignment_manager = AssignmentManager()
        if assignment_manager.save_submission(
            username=request.user.username,
            assignment_id=assignment_id,
            questions_answers=questions_answers
        ):
            return HttpResponse('File uploaded successfully')
        else:
            return HttpResponse('Failed to save submission', status=500)

    return render(request, '../static/templates/submissions/submission_portal.html')

@login_required
def submissions_home(request):
    assignment_manager = AssignmentManager()

    try:
        # Fetch assignments (you can add filters if needed)
        recent_assignments = assignment_manager.get_assignments()
        # pending_assignments = assignment_manager.get_pending_assignments()
        pending_assignments = assignment_manager.get_pending_assignments(
            username=request.user.username
        )

        context = {
            'recent_assignments': recent_assignments,
            'pending_assignments': pending_assignments,
            'total_assignments': len(recent_assignments),
            'total_pending': len(pending_assignments)
        }

        return render(request, '../static/templates/submissions/submission_portal.html', context)

    except Exception as e:
        print(f"Error in submissions home view: {e}")
        return render(request, '../static/templates/submissions/submission_portal.html', {
            'recent_assignments': [],
            'pending_assignments': [],
            'total_assignments': 0,
            'total_pending': 0,
            'error': 'Unable to fetch assignments'
        })

@login_required
def pending_submissions(request):
    assignment_manager = AssignmentManager()

    try:
        # You can add optional filtering by grade or subject
        pending_assignments = assignment_manager.get_pending_assignments(
            username=request.user.username
        )

        context = {
            'pending_assignments': pending_assignments,
            'total_pending': len(pending_assignments)
        }

        return render(request, '../static/templates/submissions/submission_pending.html', context)

    except Exception as e:
        print(f"Error in pending submissions view: {e}")
        return render(request, '../static/templates/submissions/submission_pending.html', {
            'pending_assignments': [],
            'total_pending': 0,
            'error': 'Unable to fetch pending assignments'
        })

@login_required
def submitted_assignments(request):
    assignment_manager = AssignmentManager()

    try:
        # Fetch submitted assignments (you can add filters if needed)
        submitted_assignments = assignment_manager.get_submitted_assignments()

        context = {
            'submitted_assignments': submitted_assignments,
            'total_submitted': len(submitted_assignments)
        }

        return render(request, '../static/templates/submissions/submission_submitted.html', context)

    except Exception as e:
        print(f"Error in submitted assignments view: {e}")
        return render(request, '../static/templates/submissions/submission_submitted.html', {
            'submitted_assignments': [],
            'total_submitted': 0,
            'error': 'Unable to fetch submitted assignments'
        })

from firebase_admin import credentials, firestore

db = firestore.client()
assignments_ref = db.collection('assignments')
submissions_ref = db.collection('submissions')

@login_required    
def get_submissions(request, username):
    """
    Fetch submissions for a specific user
    """
    # Ensure the user can only access their own submissions
    if username != request.user.username and not request.user.is_staff:
        return JsonResponse({"error": "Unauthorized access"}, status=403)
    
    try:
        # Create a Firestore client
        db = firestore.client()
        submissions_ref = db.collection('submissions')
        
        # Query submissions for this specific user
        submission_docs = submissions_ref.where('submitted_by', '==', username).stream()
        
        # Convert to list of dicts
        submissions = []
        for doc in submission_docs:
            submission_data = doc.to_dict()
            submission_data['id'] = doc.id  # Add the document ID
            submissions.append(submission_data)
        
        return JsonResponse({"submissions": submissions})
    
    except Exception as e:
        print(f"Error fetching submissions for {username}: {str(e)}")
        return JsonResponse({"error": str(e)}, status=400)